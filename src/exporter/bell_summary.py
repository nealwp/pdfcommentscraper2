import docx
from docx.shared import Pt, Inches
from datetime import datetime


def generate_bell_format_summary(data):

    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)
    font.bold = True

    doc.add_paragraph(f'NAME:\t\t\t\t{data["client_name"]}')
    doc.add_paragraph(f'SS#:\t\t\t\t{data["client_ssn"]}')
    doc.add_paragraph(f'TYPE OF CLAIM:\t\t{data["claimant"].claim}')
    doc.add_paragraph(f'DATE OF APPLICATION:\t{datetime.strftime(data["claimant"].pdof, "%m/%d/%Y")}')
    doc.add_paragraph(f'ALLEGED ONSET DATE:\t{data["onset_date"]}')
    doc.add_paragraph(f'DOB:\t\t\t\t{data["birthdate"]}')
    doc.add_paragraph(f'AGE:\t\t\t\t{data["age"]}')
    doc.add_paragraph(f'EDUCATION:\t\t\t{data["education"]}')
    if len(data["work_history"]) == 0:
        doc.add_paragraph('PAST WORK:\t\t\t')
    else:
        doc.add_paragraph(f'PAST WORK:\t\t\t{data["work_history"][0]["job_title"]}')
        for job in data["work_history"][1:]:
            doc.add_paragraph(f'\t\t\t\t{job["job_title"]}')

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('SEVERE IMPAIRMENTS:')

    doc.add_paragraph('')

    for idx, impairment in enumerate(data['impairments'], start=1):
        doc.add_paragraph(f'  {idx}.  {impairment}')

    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_paragraph('SUPPORTING EXHIBITS WITH TEST RESULTS:')

    doc.add_paragraph('')
    doc.add_paragraph('')

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    for exhibit in data['exhibits'].keys():

        if 'F' in exhibit:
            ex = data['exhibits'][exhibit]

            p = doc.add_paragraph()
            p.add_run(f'{exhibit}: {ex.provider_handle()}; {ex.from_date}-{ex.to_date}').bold = True
            if len(data['exhibits'][exhibit].comments) == 0:
                p = doc.add_paragraph()
                p.add_run('Not helpful')
                continue
            for comment in data['exhibits'][exhibit].comments:
                p = doc.add_paragraph()
                p.add_run(f'{comment.date}: {comment.text} ({exhibit}/{comment.exhibit_page})')

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc
