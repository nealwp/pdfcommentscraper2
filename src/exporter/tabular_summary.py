import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def generate_tablular_medical_summary(data):
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    style = doc.styles.add_style('Normal 2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(16)
    font.underline = True
    font.bold = True
    font.color.rgb = None

    doc.add_paragraph(f'Claimant:\t{data["client_name"]}')
    doc.add_paragraph(f'SSN:\t\t{data["client_ssn"]}')
    doc.add_paragraph(f'DOB:\t\t{data["birthdate"]}\t\tAge:\t{data["age"]}\tAge at AOD:\t{data["age_at_onset"]}')
    doc.add_paragraph(f'EDU:\t\t{data["education"]}')
    doc.add_paragraph(f'AOD:\t\t{data["onset_date"]}')
    doc.add_paragraph(f'PDOF:\t\t{datetime.strftime(data["claimant"].pdof, "%m/%d/%Y")}')
    doc.add_paragraph(f'Claim:\t\t{data["claimant"].claim}')
    doc.add_paragraph(f'DLI:\t\t{data["insured_date"]}\n\n')

    p = doc.add_paragraph()
    p.add_run("PAST WORK:").bold = True
    for idx, entry in enumerate(data['work_history'], start=1):
        doc.add_paragraph(f'\t{idx}. {entry["job_title"]} : {entry["intensity"]} : {entry["skill_level"]}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('DDS RFC:').bold = True
    doc.add_paragraph('Lift/carry - 99/99')
    doc.add_paragraph('Stand/walk - 99/99')
    doc.add_paragraph('Sit - 99/99')
    doc.add_paragraph('Never: ladders/ropes/scaffolds')
    doc.add_paragraph('Frequent: ramps/stairs, balance, stoop, kneel, crouch, crawl, reaching RUE')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('DDS MRFC:').bold = True
    doc.add_paragraph('something something something')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('CLAIMED MDSI PER APPLICATION:').bold = True
    doc.add_paragraph(f'{data["overview"]}\n')

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    para = doc.add_paragraph('MEDICAL EVIDENCE AND NOTES')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.style = doc.styles['Normal 2']

    for exhibit in data['exhibits'].keys():

        if 'F' in exhibit:
            ex = data['exhibits'][exhibit]
            p = doc.add_paragraph()
            p.add_run(f'{exhibit}: {ex.provider_name}; {ex.from_date}-{ex.to_date}').bold = True
            if len(data['exhibits'][exhibit].comments) == 0:
                p = doc.add_paragraph()
                p.add_run('Not helpful')
                continue
            for comment in data['exhibits'][exhibit].comments:
                p = doc.add_paragraph()
                p.add_run(f'{comment.date}: {comment.text} ({exhibit}/{comment.exhibit_page})')

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc
