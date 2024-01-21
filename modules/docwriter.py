from time import strftime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

from modules.pdfscanner import Exhibit


def generate_chronological_medical_summary(data):
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    style = doc.styles.add_style('Normal 2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(16)
    font.underline = True
    font.bold = True
    font.color.rgb = None

    doc.add_paragraph(f'CLIENT NAME\t\t\t{data["client_name"]}')
    doc.add_paragraph(f'CLIENT SSN\t\t\t{data["client_ssn"]}\n')

    doc.add_paragraph(f'TITLE\t\t\t\t{data["title"]}')
    doc.add_paragraph(f'APPLICATION DATE\t\t{data["application_date"]}')
    doc.add_paragraph(f'ALLEGED ONSET DATE\t\t{data["onset_date"]}')
    doc.add_paragraph(f'DATE LAST INSURED\t\t{data["insured_date"]}')
    doc.add_paragraph(f'PRIOR APPLICATIONS?\t\t{data["prior_applications"]}')
    doc.add_paragraph(f'DOB\t\t\t\t{data["birthdate"]}\n')

    doc.add_paragraph(f'EDUCATION\t\t\t{data["education"]}')
    doc.add_paragraph("WORK HX")
    for idx, entry in enumerate(data['work_history'], start=1):
        doc.add_paragraph(f'\t{idx}. {entry["job_title"]} : {entry["intensity"]} : {entry["skill_level"]}')

    doc.add_paragraph('')
    doc.add_paragraph(f'DRUG USE\t\t\t{data["drug_use"]}')
    doc.add_paragraph(f'CRIMINAL HX\t\t\t{data["criminal_history"]}\n')

    doc.add_paragraph(f'{data["overview"]}\n')

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    para = doc.add_paragraph('MEDICAL EVIDENCE AND NOTES')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.style = doc.styles['Normal 2']

    for comment in data['comments']:
        p = doc.add_paragraph()
        p.add_run(f'{datetime.strftime(comment["date"], "%m/%d/%Y")}: ').bold = True
        p.add_run(f'{comment["text"]}')
        p.add_run(f' [Exhibit: {comment["ref"]}]').bold = True

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc


def generate_tablular_medical_summary(data):
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    style = doc.styles.add_style('Normal 2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(16)
    font.underline = True
    font.bold = True
    font.color.rgb = None

    # p = doc.add_paragraph()
    #p.add_run(f'I think this case is a total winner because \
    #    everything is perfect and there\'s literally no reason \
    #    why anyone would say otherwise.').bold = True
    # p = doc.add_paragraph()
    #r = p.add_run('Case Rating: 5/5')
    #r.bold = True
    #r.font.size = Pt(24)
    #p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph(f'Claimant:\t{data["client_name"]}')
    doc.add_paragraph(f'SSN:\t\t{data["client_ssn"]}')
    doc.add_paragraph(f'DOB:\t\t{data["birthdate"]}\t\tAge:\t{data["age"]}\tAge at AOD:\t{data["age_at_onset"]}')
    doc.add_paragraph(f'EDU:\t\t{data["education"]}')
    doc.add_paragraph(f'AOD:\t\t{data["onset_date"]}')
    doc.add_paragraph(f'PDOF:\t\t{datetime.strftime(data["claimant"].pdof, "%m/%d/%Y")}')
    doc.add_paragraph(f'Claim:\t\t{data["claimant"].claim}')
    doc.add_paragraph(f'DLI:\t\t{data["insured_date"]}\n\n')
    
    p = doc.add_paragraph()
    p.add_run("PAST WORK:").bold = True
    for idx, entry in enumerate(data['work_history'], start=1):
        doc.add_paragraph(f'\t{idx}. {entry["job_title"]} : {entry["intensity"]} : {entry["skill_level"]}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('DDS RFC:').bold = True
    doc.add_paragraph(f'Lift/carry - 99/99')
    doc.add_paragraph(f'Stand/walk - 99/99')
    doc.add_paragraph(f'Sit - 99/99')
    doc.add_paragraph(f'Never: ladders/ropes/scaffolds')
    doc.add_paragraph(f'Frequent: ramps/stairs, balance, stoop, kneel, crouch, crawl, reaching RUE')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('DDS MRFC:').bold = True
    doc.add_paragraph(f'something something something')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('CLAIMED MDSI PER APPLICATION:').bold = True
    doc.add_paragraph(f'{data["overview"]}\n')
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    para = doc.add_paragraph('MEDICAL EVIDENCE AND NOTES')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.style = doc.styles['Normal 2']

    for exhibit in data['exhibits'].keys():
        
        if 'F' in exhibit:
            ex = data['exhibits'][exhibit]
            p = doc.add_paragraph()
            p.add_run(f'{exhibit}: {ex.provider_name}; {ex.from_date}-{ex.to_date}').bold = True
            if len(data['exhibits'][exhibit].comments) == 0:
                p = doc.add_paragraph()
                p.add_run('Not helpful')
                continue
            for comment in data['exhibits'][exhibit].comments:
                p = doc.add_paragraph()
                p.add_run(f'{comment.date}: {comment.text} ({exhibit}/{comment.exhibit_page})')

    #for comment in data['comments']:
    #    p = doc.add_paragraph()
    #    p.add_run(f'{datetime.strftime(comment.date, "%m/%d/%Y")}: ').bold = True
    #    p.add_run(f'{comment.text}')
    #    p.add_run(f' ({str(comment["ref"]).replace("-","/")})')

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc


def generate_tablular_medical_summary_v2(data):

    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)
    font.bold = True

    doc.add_paragraph(f'NAME:\t\t\t\t{data["client_name"]}')
    doc.add_paragraph(f'SS#:\t\t\t\t{data["client_ssn"]}')
    doc.add_paragraph(f'TYPE OF CLAIM:\t\t{data["claimant"].claim}')
    doc.add_paragraph(f'DATE OF APPLICATION:\t{datetime.strftime(data["claimant"].pdof, "%m/%d/%Y")}')
    doc.add_paragraph(f'ALLEGED ONSET DATE:\t{data["onset_date"]}')
    doc.add_paragraph(f'DOB:\t\t\t\t{data["birthdate"]}')
    doc.add_paragraph(f'AGE:\t\t\t\t{data["age"]}')
    doc.add_paragraph(f'EDUCATION:\t\t\t{data["education"]}')
    doc.add_paragraph(f'PAST WORK:\t\t\t{data["work_history"][0]["job_title"]}')
    doc.add_paragraph(f'\t\t\t\t{data["work_history"][1]["job_title"]}')
    doc.add_paragraph(f'\t\t\t\t{data["work_history"][2]["job_title"]}')
    doc.add_paragraph(f'\t\t\t\t{data["work_history"][3]["job_title"]}')

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('SEVERE IMPAIRMENTS:')

    doc.add_paragraph('')

    for idx, impairment in enumerate(data['impairments'], start=1):
        doc.add_paragraph(f'  {idx}.  {impairment}')

    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_paragraph('SUPPORING EXHIBITS WITH TEST RESULTS:')

    doc.add_paragraph('')
    doc.add_paragraph('')

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    for exhibit in data['exhibits'].keys():

        if 'F' in exhibit:
            ex = data['exhibits'][exhibit]
            p = doc.add_paragraph()
            p.add_run(f'{exhibit}: {ex.provider_name}; {ex.from_date}-{ex.to_date}').bold = True
            if len(data['exhibits'][exhibit].comments) == 0:
                p = doc.add_paragraph()
                p.add_run('Not helpful')
                continue
            for comment in data['exhibits'][exhibit].comments:
                p = doc.add_paragraph()
                p.add_run(f'{comment.date}: {comment.text} ({exhibit}/{comment.exhibit_page})')

    #for comment in data['comments']:
    #    p = doc.add_paragraph()
    #    p.add_run(f'{datetime.strftime(comment.date, "%m/%d/%Y")}: ').bold = True
    #    p.add_run(f'{comment.text}')
    #    p.add_run(f' ({str(comment["ref"]).replace("-","/")})')

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc


def main():
    exhibits = {'1A': Exhibit(provider_name='Disability Determination Explanation - DDE Dec. Dt.', from_date='08/18/2021', to_date='08/18/2021'), '2A': Exhibit(provider_name='Disability Determination Transmittal - 831 Dec. Dt.', from_date='08/18/2021', to_date='08/18/2021'), '3A': Exhibit(provider_name='Disability Determination Explanation - DDE Dec. Dt.', from_date='01/05/2022', to_date='01/05/2022'), '4A': Exhibit(provider_name='Disability Determination Transmittal - 831 Dec. Dt.', from_date='01/05/2022', to_date='01/05/2022'), '1B': Exhibit(provider_name='T16 Notice of Disapproved Claim - L444', from_date='08/18/2021', to_date='08/18/2021'), '2B': Exhibit(provider_name='Fee Agreement for Representation before SSA - FEEAGRMT', from_date='10/05/2021', to_date='10/05/2021'), '3B': Exhibit(provider_name="SSA-1696 - Claimant's Appointment of a Representative - 1696", from_date='10/25/2021', to_date='10/25/2021'), '4B': Exhibit(provider_name='Request for Reconsideration - 561', from_date='10/25/2021', to_date='10/25/2021'), '5B': Exhibit(provider_name='Statement of Good Cause for Untimely Filing - GDCAUSE', from_date='11/02/2021', to_date='11/02/2021'), '6B': Exhibit(provider_name='T16 Disability Reconsideration Notice - L1130', from_date='01/06/2022', to_date='01/06/2022'), '7B': Exhibit(provider_name='Request for Hearing by ALJ - 501', from_date='01/20/2022', to_date='01/20/2022'), '8B': Exhibit(provider_name='COVID Hearing Agreement Form - COVIDAGREEFRM', from_date='01/24/2022', to_date='01/24/2022'), '9B': Exhibit(provider_name='Request for Hearing Acknowledgement Letter - HRGACK', from_date='02/03/2022', to_date='02/03/2022'), '10B': Exhibit(provider_name='COVID Hearing Agreement Form - COVIDAGREEFRM', from_date='02/11/2022', to_date='02/11/2022'), '11B': Exhibit(provider_name='Hearing Notice - 507', from_date='05/24/2022', to_date='05/24/2022'), '12B': Exhibit(provider_name='COVID Hearing Agreement Form - COVIDAGREEFRM', from_date='07/06/2022', to_date='07/06/2022'), '1D': Exhibit(provider_name='Application for Supplemental Security Income Benefits - 8000', from_date='03/15/2021', to_date='03/15/2021'), '2D': Exhibit(provider_name='Misc Non-Disability Development - MDF D', from_date='03/15/2021', to_date='03/15/2021'), '3D': Exhibit(provider_name='Detailed Earnings Query - DEQY', from_date='07/06/2022', to_date='07/06/2022'), '4D': Exhibit(provider_name='New Hire, Quarter Wage, Unemployment Query (NDNH) - NDNH', from_date='07/06/2022', to_date='07/06/2022'), '5D': Exhibit(provider_name='Summary Earnings Query - SEQY', from_date='07/06/2022', to_date='07/06/2022'), '6D': Exhibit(provider_name='Certified Earnings Records - CERTERN', from_date='07/06/2022', to_date='07/06/2022'), '7D': Exhibit(provider_name='DISCO DIB Insured Status Report - DISCODIB', from_date='07/06/2022', to_date='07/06/2022'), '1E': Exhibit(provider_name='Disability Report - Field Office - 3367', from_date='03/16/2021', to_date='03/16/2021'), '2E': Exhibit(provider_name='Disability Report - Adult - 3368 Claimant', from_date='03/16/2021', to_date='03/16/2021'), '3E': Exhibit(provider_name='Seizure Questionnaire - SZREQ Claimant', from_date='05/04/2021', to_date='05/04/2021'), '4E': Exhibit(provider_name='Exertional Activities Questionnaire - EXRTNLQ Claimant', from_date='05/04/2021', to_date='05/04/2021'), '5E': Exhibit(provider_name='Work History Report - 3369 Claimant', from_date='05/12/2021', to_date='05/12/2021'), '6E': Exhibit(provider_name='Case Development Worksheet - CDW DDS', from_date='03/17/2021', to_date='08/18/2021'), '7E': Exhibit(provider_name='Disability Report - Field Office - 3367', from_date='11/19/2021', to_date='11/19/2021'), '8E': Exhibit(provider_name='Disability Report - Appeals - 3441 Representative', from_date='11/19/2021', to_date='11/19/2021'), '9E': Exhibit(provider_name='Case Development Worksheet - CDW DDS', from_date='11/19/2021', to_date='01/06/2022'), '10E': Exhibit(provider_name='Disability Report - Appeals - 3441 Representative', from_date='01/22/2022', to_date='01/22/2022'), '11E': Exhibit(provider_name='Disability Report - Field Office - 3367', from_date='01/22/2022', to_date='01/22/2022'), '12E': Exhibit(provider_name='Correspondence regarding efforts to obtain evidence - EALTR Representative', from_date='06/14/2022', to_date='06/14/2022'), '13E': Exhibit(provider_name='Exhibit List to Rep PH2E - EXHIBITLISTREP OHO', from_date='07/11/2022', to_date='07/11/2022'), '1F': Exhibit(provider_name='Progress Notes - PROGRESSNOTES COPPER RIDGE CARE CENTER–Mantri MD', from_date='11/18/2019', to_date='11/20/2019'), '2F': Exhibit(provider_name='HIT MER - HITMER Dignity Health–A Ayyad MD', from_date='12/18/2019', to_date='12/18/2019'), '3F': Exhibit(provider_name='Office Treatment Records - OFFCREC Dignity Health–North State–Ayyad MD', from_date='12/18/2019', to_date='12/18/2019'), '4F': Exhibit(provider_name='Office Treatment Records - OFFCREC TRINITY COMMUNITY HEALTH CLN', from_date='04/06/2020', to_date='06/30/2020'), '5F': Exhibit(provider_name='Office Treatment Records - OFFCREC HARVINDER S BIRK MD', from_date='06/23/2020', to_date='10/20/2020'), '6F': Exhibit(provider_name='Claimant-supplied Evidence - CLMTEVID comingled sources', from_date='10/30/2016', to_date='10/20/2020'), '7F': Exhibit(provider_name='Physical/Occupational Therapy Records - PTREC Mountain Physical Therapy–M Thomas DPT', from_date='06/23/2021', to_date='06/23/2021'), '8F': Exhibit(provider_name='Medical Report/General - 3826 MOUNTAIN VALLEY PHYSICAL THRPY–M Thomas DPT', from_date='06/23/2021', to_date='06/23/2021'), '9F': Exhibit(provider_name='Medical Report/General - 3826 TRINITY COMMUNITY HEALTH CLN–Jared Fife PA-C', from_date='07/01/2021', to_date='07/01/2021'), '10F': Exhibit(provider_name='Treating Source Statement - TREATSRCSTATE Mountain Physical Therapy–Melissa Thomas DPT', from_date='07/21/2021', to_date='07/21/2021'), '11F': Exhibit(provider_name='CE Orthopedic - CEORTHO Dale H Van Kirk MD', from_date='08/01/2021', to_date='08/01/2021'), '12F': Exhibit(provider_name='HIT MER - HITMER Dignity Health #2', from_date='08/11/2021', to_date='08/11/2021'), '13F': Exhibit(provider_name='Physical/Occupational Therapy Records - PTREC MOUNTAIN VALLEY PHYSICAL THRPY–M Thomas DPT', from_date='02/14/2020', to_date='10/12/2021'), '14F': Exhibit(provider_name='Outpatient Hospital Records - OUTHOSP TRINITY COMMUNITY HEALTH CLN', from_date='05/11/2021', to_date='04/21/2022'), '15F': Exhibit(provider_name='HIT MER - HITMER Dignity Health #3–J Campanelli M', from_date='01/17/2022', to_date='01/17/2022')}
    print(exhibits.keys())
    for exhibit in exhibits.keys():
        ex = exhibits[exhibit]
        print(f'{exhibit}: {ex.provider_name}; {ex.from_date}-{ex.to_date}')
    return


if __name__ == "__main__":
    main()
