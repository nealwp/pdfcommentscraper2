from time import strftime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def generate_medical_summary(data):
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    style = doc.styles.add_style('Normal 2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(16)
    font.underline = True
    font.bold = True
    font.color.rgb = None 

    doc.add_paragraph(f'CLIENT NAME\t\t\t{data["client_name"]}')
    doc.add_paragraph(f'CLIENT SSN\t\t\t{data["client_ssn"]}\n')

    doc.add_paragraph(f'TITLE\t\t\t\t{data["title"]}')
    doc.add_paragraph(f'APPLICATION DATE\t\t{data["application_date"]}')
    doc.add_paragraph(f'ALLEGED ONSET DATE\t\t{data["onset_date"]}')
    doc.add_paragraph(f'DATE LAST INSURED\t\t{data["insured_date"]}')
    doc.add_paragraph(f'PRIOR APPLICATIONS?\t\t{data["prior_applications"]}')
    doc.add_paragraph(f'DOB\t\t\t\t{data["birthdate"]}\n')

    doc.add_paragraph(f'EDUCATION\t\t\t{data["education"]}')
    doc.add_paragraph("WORK HX")
    for idx, entry in enumerate(data['work_history'], start=1):
        doc.add_paragraph(f'\t{idx}. {entry["job_title"]} : {entry["intensity"]} : {entry["skill_level"]}')

    doc.add_paragraph('')
    doc.add_paragraph(f'DRUG USE\t\t\t{data["drug_use"]}')
    doc.add_paragraph(f'CRIMINAL HX\t\t\t{data["criminal_history"]}\n')

    doc.add_paragraph(f'{data["overview"]}\n')
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = doc.styles['Normal']

    para = doc.add_paragraph('MEDICAL EVIDENCE AND NOTES')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.style = doc.styles['Normal 2']

    for comment in data['comments']:
        p = doc.add_paragraph()
        p.add_run(f'{datetime.strftime(comment["date"], "%m/%d/%Y")}: ').bold = True
        p.add_run(f'{comment["text"]}')
        p.add_run(f' [Exhibit: {comment["ref"]}]').bold = True

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    return doc
