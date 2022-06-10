from pydoc import cli
import re
from time import perf_counter
from csv import writer, DictWriter
from tkinter import BOTH, END, LEFT, LabelFrame, Listbox, Scrollbar, StringVar, Text, Tk, Button, Label, Entry, Toplevel, Menu, Frame, SUNKEN, N, S, E, W, BOTTOM, X
from tkinter.ttk import Combobox
from tkcalendar import Calendar, DateEntry
from commentscraper import scrape_comments
from pdfscanner import scanforkeywords, scan_for_client_info
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
import json


from helpers import write_csv, get_file_path, get_save_path

class App(Tk):
    def __init__(self, *args, **kwargs) -> None:
        super().__init__()
        self.attributes('-alpha', 0.0)
        self.geometry('809x500')
        self.title('disabilitydude')

        self.menubar = Menu(self, bd=5)

        self.filemenu = Menu(self.menubar, tearoff=0)
        self.filemenu.add_command(label='New Summary', command=self.open_new_summary_form)
        self.filemenu.add_command(label='Open Summary')
        self.filemenu.add_command(label="Exit", command=self.destroy)
        self.menubar.add_cascade(label="File", menu=self.filemenu)

        self.actions_menu = Menu(self.menubar, tearoff=0)
        self.actions_menu.add_command(label="Scan PDF for Keywords", command=self.run_keyword_scan)
        self.actions_menu.add_command(label="Scrape Comments from PDF", command=run_commentscraper)
        self.menubar.add_cascade(label="Actions", menu=self.actions_menu)
        
        self.settingsmenu = Menu(self.menubar, tearoff=0)
        self.settingsmenu.add_command(label="Keywords", command=self.open_keywords_menu)
        self.settingsmenu.add_command(label="Context Size", command=self.open_context_size_menu)
        self.menubar.add_cascade(label="Settings", menu=self.settingsmenu)
        
        self.menubar.add_command(label="Help", command=lambda: self.set_status('you clicked help'))
        
        self.config(menu=self.menubar)
        
        self.status_bar = Frame(self, bd=1, relief=SUNKEN)
        
        self.status = Label(self.status_bar, text='Ready', anchor=W, padx=5)
        self.status.grid(column=0, row=0, sticky='w', columnspan=1)
        
        self.version = Label(self.status_bar, text='v1.0.0', anchor=E, padx=5)
        self.version.grid(column=1, row=0, sticky="e", columnspan=1)
        
        self.status_bar.columnconfigure(0, weight=1)
        self.status_bar.pack(side=BOTTOM, fill=X)
        
        self._center()
        self.attributes('-alpha', 1.0)

    def set_status(self, message):
        self.status = Label(self.status_bar, text=message, anchor=W, padx=5)
        self.status.grid(column=0, row=0, sticky='w', columnspan=1)

    def run_keyword_scan(self):

        fieldnames = ['keyword', 'page', 'text', 'exhibit', 'provider', 'exhibit_page']

        pdf_path = get_file_path()
        debut = perf_counter()
        output = scanforkeywords(pdf_path, self)
        fin = perf_counter()
        output_path = get_save_path()
        with open(output_path,'w',encoding='utf-8', newline='') as f:
            csvdw = DictWriter(f, fieldnames=fieldnames)
            csvdw.writeheader()
            csvdw.writerows(output)
        print(f"Completed in {fin - debut:0.4f}s")
        return

    def open_keywords_menu(self):

        keywords = read_keywords_file()

        win = Toplevel(self)
        win.attributes('-alpha', 0.0)
        win.title('Keywords Settings')

        leftFrame = Frame(win)
        leftFrame.grid(column=0, row=1)

        entrylabel = Label(leftFrame, text='New Keyword')
        entrylabel.grid(column=0, row=0, sticky='sw', rowspan=1, padx=10)

        entrybox = Entry(leftFrame, width=24, bd=1)
        entrybox.grid(column=0, row=1, sticky='nw', rowspan=1, columnspan=1, padx=10)

        addBtn = Button(leftFrame, text='Add', command=lambda: add_keyword(entrybox.get(), entrybox))
        addBtn.grid(column=0, row=2, pady=10)
        addBtn.bind("<Enter>", on_enter)
        addBtn.bind("<Leave>", on_leave)

        rightFrame = Frame(win)
        rightFrame.grid(column=1, row=1, pady=10)
        
        keywordboxlabel = Label(rightFrame, text='Active Keywords')
        keywordboxlabel.grid(column=0, row=0, sticky='nw', rowspan=1)

        listbox = Listbox(rightFrame, height=16, selectmode='single', activestyle='none')
        for i, kw in enumerate(keywords):
            listbox.insert(i, kw)

        listbox.grid(column=0, row=1, sticky='nw')

        scrollbar = Scrollbar(rightFrame, orient='vertical', command=listbox.yview)
        scrollbar.grid(column=1, row=1, sticky='nsw')
        listbox['yscrollcommand'] = scrollbar.set

        deleteBtn = Button(rightFrame, text='Delete Selected', command=lambda: delete_keyword(listbox.get(listbox.curselection())))
        deleteBtn.grid(column=0, row=2, pady=10)
        deleteBtn.bind("<Enter>", on_enter)
        deleteBtn.bind("<Leave>", on_leave)
        
        doneBtn = Button(win, text='Done', command=win.destroy)
        doneBtn.bind("<Enter>", on_enter)
        doneBtn.bind("<Leave>", on_leave)
        doneBtn.grid(column=0, row=3, columnspan=2, pady=10)
        center(win)
        win.attributes('-alpha', 1.0)
        win.mainloop()

    def open_context_size_menu(self):
        win = Toplevel(self, padx=15, pady=15)
        win.attributes('-alpha', 0.0)
        win.title('Context Size Settings')

        words_minus_label = Label(win, text="Words Before Keyword:")
        words_minus_label.grid(column=0, row=0, sticky='w', padx=10)
        
        words_minus_entry = Entry(win, width=3, bd=1)
        words_minus_entry.grid(column=1, row=0, sticky='e', padx=10)
        
        words_plus_label = Label(win, text="Words After Keyword:")
        words_plus_label.grid(column=0, row=1, sticky='w', padx=10)
        
        words_plus_entry = Entry(win, width=3, bd=1)
        words_plus_entry.grid(column=1, row=1, sticky='e', padx=10)

        done_btn = Button(win, text='Done', command=win.destroy)
        done_btn.bind("<Enter>", on_enter)
        done_btn.bind("<Leave>", on_leave)
        done_btn.grid(column=0, row=3, columnspan=2)
        
        center(win)
        win.attributes('-alpha', 1.0)
        win.mainloop()
    
    def open_new_summary_form(self):
        self.summary_form = SummaryForm(self)

    def _center(self):
        """
        centers a tkinter window
        :param win: the main window or Toplevel window to center
        """
        self.update_idletasks()
        width = self.winfo_width()
        frm_width = self.winfo_rootx() - self.winfo_x()
        win_width = width + 2 * frm_width
        height = self.winfo_height()
        titlebar_height = self.winfo_rooty() - self.winfo_y()
        win_height = height + titlebar_height + frm_width
        x = self.winfo_screenwidth() // 2 - win_width // 2
        y = self.winfo_screenheight() // 2 - win_height // 2
        self.geometry('{}x{}+{}+{}'.format(width, height, x, y))
        self.deiconify()

class SummaryForm(Toplevel):
    def __init__(self, parent):
        super().__init__(parent, padx=15, pady=15)
        self.parent = parent
        self.work_history = []
        self.attributes('-alpha', 0.0)
        self.title('New Summary')
        
        self.detect_client_btn = Button(self, text="Detect Client Info", command=self._detect_client_info)
        self.detect_client_btn.bind("<Enter>", on_enter)
        self.detect_client_btn.bind("<Leave>", on_leave)
        self.detect_client_btn.grid(column=0, row=0, columnspan=2, sticky='w', padx=10, pady=2)

        self.name_label = Label(self, text="Client Name:")
        self.name_label.grid(column=0, row=1, sticky='w', padx=10, pady=2)
        
        self.name_entry = Entry(self, bd=1, width=32)
        self.name_entry.grid(column=1, row=1, sticky='e', padx=10, pady=2)
        
        self.ssn_label = Label(self, text="SSN:")
        self.ssn_label.grid(column=0, row=2, sticky='w', padx=10, pady=2)
        
        self.ssn_entry = Entry(self, bd=1, width=32)
        self.ssn_entry.grid(column=1, row=2, sticky='e', padx=10, pady=2)

        self.title_label = Label(self, text="Title:")
        self.title_label.grid(column=0, row=3, sticky='w', padx=10, pady=2)
        
        self.title_entry = Entry(self, bd=1, width=32)
        self.title_entry.grid(column=1, row=3, sticky='e', padx=10, pady=2)

        self.application_date_label = Label(self, text="Application Date:")
        self.application_date_label.grid(column=0, row=4, sticky='w', padx=10, pady=2)

        #application_date_entry = Calendar(win, selectmode='day')
        self.application_date_entry = Entry(self, bd=1, width=32)
        self.application_date_entry.grid(column=1, row=4, sticky='e', padx=10, pady=2)

        self.onset_date_label = Label(self, text="Alleged Onset Date:")
        self.onset_date_label.grid(column=0, row=5, sticky='w', padx=10, pady=2)

        self.onset_date_entry = Entry(self, bd=1, width=32)
        self.onset_date_entry.grid(column=1, row=5, sticky='e', padx=10, pady=2)

        self.insured_date_label = Label(self, text="Last Insured Date:")
        self.insured_date_label.grid(column=0, row=6, sticky='w', padx=10, pady=2)

        self.insured_date_entry = Entry(self, bd=1, width=32)
        self.insured_date_entry.grid(column=1, row=6, sticky='e', padx=10, pady=2)

        self.prior_label = Label(self, text="Prior Applications?")
        self.prior_label.grid(column=0, row=7, sticky='w', padx=10, pady=2)

        self.prior_txtVar = StringVar()
        self.prior_combo = Combobox(self, width=6, textvariable=self.prior_txtVar, justify='left')
        self.prior_combo['values'] = ('No', 'Yes')
        self.prior_combo.grid(column=1, row=7, sticky='e', padx=10, pady=2)
        self.prior_combo.current(0)

        self.birthday_label = Label(self, text="Date of Birth:")
        self.birthday_label.grid(column=0, row=8, sticky='w', padx=10, pady=2)

        self.birthday_entry = Entry(self, bd=1, width=32)
        self.birthday_entry.grid(column=1, row=8, sticky='e', padx=10, pady=2)

        self.education_label = Label(self, text="Education:")
        self.education_label.grid(column=0, row=9, sticky='w', padx=10, pady=2)

        self.education_entry = Entry(self, bd=1, width=32)
        self.education_entry.grid(column=1, row=9, sticky='e', padx=10, pady=2)

        self.drugs_label = Label(self, text="Drug Use?")
        self.drugs_label.grid(column=0, row=10, sticky='w', padx=10, pady=2)

        self.drugs_txtVar = StringVar()
        self.drugs_combo = Combobox(self, width=6, textvariable=self.drugs_txtVar, justify='left')
        self.drugs_combo['values'] = ('No', 'Yes')
        self.drugs_combo.grid(column=1, row=10, sticky='e', padx=10, pady=2)
        self.drugs_combo.current(0)

        self.criminal_label = Label(self, text="Criminal History?")
        self.criminal_label.grid(column=0, row=11, sticky='w', padx=10, pady=2)

        self.criminal_txtVar = StringVar()
        self.criminal_combo = Combobox(self, width=6, textvariable=self.criminal_txtVar, justify='left')
        self.criminal_combo['values'] = ('No', 'Yes')
        self.criminal_combo.grid(column=1, row=11, sticky='e', padx=10, pady=2)
        self.criminal_combo.current(0)

        self.overview_label = Label(self, text="Case Overview:")
        self.overview_label.grid(column=0, row=12, sticky='w', padx=10, pady=2)

        self.overview_text = Text(self, height=6, font='Calibri', wrap='word')
        self.overview_text.grid(column=0, row=13, columnspan=4, rowspan=2, padx=10, pady=2)

        self.add_work_history_btn = Button(self, text="Add Work History", command=self._open_work_history_form)
        self.add_work_history_btn.bind("<Enter>", on_enter)
        self.add_work_history_btn.bind("<Leave>", on_leave)
        self.add_work_history_btn.grid(column=2, row=0, columnspan=2)

        
        self.work_history_frame = LabelFrame(self, text="Work History")
        self.work_history_frame.grid(column=2, row=1, columnspan=2, rowspan=12, sticky='n', pady=10)

        done_btn = Button(self, text='Done', command=self._save_to_file)
        done_btn.bind("<Enter>", on_enter)
        done_btn.bind("<Leave>", on_leave)
        done_btn.grid(column=0, row=99, columnspan=4, pady=4)

        center(self)
        self.attributes('-alpha', 1.0)
        self.mainloop()

    def _open_work_history_form(self):
        self.work_history_form = WorkHistoryForm(self)

    def _detect_client_info(self):
        pdf_path = get_file_path()
        client_info = scan_for_client_info(pdf_path)
        self.name_entry.insert(0, client_info['Claimant'])
        self.ssn_entry.insert(0, client_info['SSN'])
        self.onset_date_entry.insert(0, client_info['Alleged Onset'])
        self.title_entry.insert(0, client_info['Claim Type'])
        self.application_date_entry.insert(0, client_info['Application'])
        self.update()
        self.deiconify()

    def _paint_work_history(self):
        for i, entry in enumerate(self.work_history, start=1):
            entry_text = f'{i}. {entry["job_title"]} : {entry["intensity"]} : {entry["skill_level"]}'
            label = Label(self.work_history_frame, text=entry_text)
            label.grid(column=0, row=i+1, padx=10, pady=2, sticky='w')

    def _save_to_file(self):
        
        client_name = self.name_entry.get()
        client_ssn = self.ssn_entry.get()
        title = self.title_entry.get()
        application_date = self.application_date_entry.get()
        onset_date = self.onset_date_entry.get()
        insured_date = self.insured_date_entry.get()
        prior_applications = self.prior_txtVar.get()
        birthdate = self.birthday_entry.get()
        education = self.education_entry.get()
        work_history = []
        drug_use = self.drugs_txtVar.get()
        criminal_history = self.criminal_txtVar.get()
        overview = self.overview_text.get('1.0', 'end-1c')

        data = {
            'client_name': client_name,
            'client_ssn': client_ssn,
            'title': title,
            'application_date': application_date,
            'onset_date': onset_date,
            'insured_date': insured_date,
            'prior_applications': prior_applications,
            'birthdate': birthdate,
            'education': education,
            'work_history': work_history,
            'drug_use': drug_use,
            'criminal_history': criminal_history,
            'overview': overview
        }

        print(data)

        doc = docx.Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'

        style = doc.styles.add_style('Normal 2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Arial Narrow'
        font.size = Pt(16)
        font.underline = True
        font.bold = True
        font.color.rgb = None 

        doc.add_paragraph(f'CLIENT NAME\t\t\t{data["client_name"]}')
        doc.add_paragraph(f'CLIENT SSN\t\t\t{data["client_ssn"]}\n')
 
        doc.add_paragraph(f'TITLE\t\t\t\t{data["title"]}')
        doc.add_paragraph(f'APPLICATION DATE\t\t{data["application_date"]}')
        doc.add_paragraph(f'ALLEGED ONSET DATE\t\t{data["onset_date"]}')
        doc.add_paragraph(f'DATE LAST INSURED\t\t{data["insured_date"]}')
        doc.add_paragraph(f'PRIOR APPLICATIONS?\t\t{data["prior_applications"]}')
        doc.add_paragraph(f'DOB\t\t\t\t{data["birthdate"]}\n')

        doc.add_paragraph(f'EDUCATION\t\t\t{data["education"]}')
        doc.add_paragraph("WORK HX")
 
        doc.add_paragraph(f'DRUG USE\t\t\t{data["drug_use"]}')
        doc.add_paragraph(f'CRIMINAL HX\t\t\t{data["criminal_history"]}\n')

 
        doc.add_paragraph(f'{data["overview"]}\n')
        
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.style = doc.styles['Normal']

        para = doc.add_paragraph('MEDICAL EVIDENCE AND NOTES')
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.style = doc.styles['Normal 2']

        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        output_path = get_save_path('docx')
        doc.save(output_path)

        self.parent.set_status('Data Saved')
        self.destroy()
        return

class WorkHistoryForm(Toplevel):
    def __init__(self, parent):
        super().__init__(parent, padx=15, pady=15)
        self.parent = parent
        self.attributes('-alpha', 0.0)
        self.title('Work History')

        self.job_title_label = Label(self, text="Job Title")
        self.job_title_label.grid(column=0, row=0, padx=10, pady=2, sticky='w')

        self.job_title_entry = Entry(self, bd=1)
        self.job_title_entry.grid(column=0, row=1, padx=10, pady=2)

        self.intensity_label = Label(self, text="Intensity")
        self.intensity_label.grid(column=1, row=0, padx=10, pady=2, sticky='w')

        self.intensity_txtVar = StringVar()
        self.intensity_combo = Combobox(self, textvariable=self.intensity_txtVar, justify='left')
        self.intensity_combo['values'] = ('Light', 'Medium', 'Heavy')
        self.intensity_combo.grid(column=1, row=1, padx=10, pady=2)
        self.intensity_combo.current(0)

        self.skill_level_label = Label(self, text="Skill Level")
        self.skill_level_label.grid(column=2, row=0, padx=10, pady=2, sticky='w')

        self.skill_level_txtVar = StringVar()
        self.skill_level_combo = Combobox(self, textvariable=self.skill_level_txtVar, justify='left')
        self.skill_level_combo['values'] = ('Unskilled', 'Semi-Skilled', 'Skilled')
        self.skill_level_combo.grid(column=2, row=1, padx=10, pady=2)
        self.skill_level_combo.current(0)

        self.add_entry_btn = Button(self, text='Add', command=self._add_entry)
        self.add_entry_btn.grid(column=3, row=1, padx=10, pady=2)

        center(self)
        self.attributes('-alpha', 1.0)
        self.mainloop()
    
    def _add_entry(self):
        entry = {
            'job_title': self.job_title_entry.get(),
            'intensity': self.intensity_txtVar.get(),
            'skill_level': self.skill_level_txtVar.get()
        }

        self.parent.work_history.append(entry)
        self.parent._paint_work_history()
        self.parent.update()

class StatusBar(App):
    
    def __init__(self):
        super().__init__()
        self.parent = parent
        self.parent.status_bar = Frame(self.parent, bd=1, relief=SUNKEN)
        
        self.status = Label(self.parent.status_bar, text='Ready', anchor=W, padx=5)
        self.status.grid(column=0, row=0, sticky='w', columnspan=1)
        
        self.version = Label(self.parent.status_bar, text='v1.0.0', anchor=E, padx=5)
        self.version.grid(column=1, row=0, sticky="e", columnspan=1)
        
        self.parent.status_bar.columnconfigure(0, weight=1)
        self.parent.status_bar.pack(side=BOTTOM, fill=X)

    def set_status(self, message):
        self.status = Label(self.parent.status_bar, text=message, anchor=W, padx=5)

class MenuBar(App):
    def __init__(self) -> None:
        super().__init__()
        self.parent = parent
        self.parent.menubar = Menu(self.parent, bd=5)

        self.filemenu = Menu(self.parent.menubar, tearoff=0)
        self.filemenu.add_command(label="Exit", command=parent.destroy)
        self.parent.menubar.add_cascade(label="File", menu=self.filemenu)

        self.actions_menu = Menu(self.parent.menubar, tearoff=0)
        self.actions_menu.add_command(label="Scan PDF for Keywords", command=run_keyword_scan)
        self.actions_menu.add_command(label="Scrape Comments from PDF", command=run_commentscraper)
        self.parent.menubar.add_cascade(label="Actions", menu=self.actions_menu)
        
        self.settingsmenu = Menu(self.parent.menubar, tearoff=0)
        self.settingsmenu.add_command(label="Keywords", command=lambda: open_keywords_menu(parent))
        self.settingsmenu.add_command(label="Context Size", command=lambda: open_context_size_menu(parent))
        self.parent.menubar.add_cascade(label="Settings", menu=self.settingsmenu)
        
        self.parent.menubar.add_command(label="Help", command=lambda: self.set_status('you clicked help'))
        
        self.parent.config(menu=self.parent.menubar)

    def set_app_status(self, message):
        self.set_status(message)


def findWholeWord(w):
    return re.compile(r'\b^({0})$\b'.format(w), flags=re.IGNORECASE).search

def run_keyword_scan():

    fieldnames = ['keyword', 'page', 'text', 'exhibit', 'provider', 'exhibit_page']

    pdf_path = get_file_path()
    debut = perf_counter()
    output = scanforkeywords(pdf_path)
    fin = perf_counter()
    output_path = get_save_path()
    with open(output_path,'w',encoding='utf-8', newline='') as f:
        csvdw = DictWriter(f, fieldnames=fieldnames)
        csvdw.writeheader()
        csvdw.writerows(output)
    print(f"Completed in {fin - debut:0.4f}s")
    return

def read_keywords_file():
    with open(r".\\config\\keywords", "r") as keyword_file:
        file_content = keyword_file.read()
        keywords = file_content.split('\n')
        keywords = [w for w in keywords if w]
        keywords.sort()
    return keywords

def add_keyword(word, entrybox):
    entrybox.delete(0, END)
    with open(r".\\config\\keywords", "a") as keyword_file:
        keyword_file.write(word + '\n')
    return

def delete_keyword(word):
    keywords = read_keywords_file()
    for i, kw in enumerate(keywords):
        if kw == word:
            keywords.pop(i)
    
    with open(r".\\config\\keywords", "w") as keyword_file:
        keyword_file.writelines(kw + '\n' for kw in keywords)
    return

def run_commentscraper():
    scrape_comments()
    return

def on_enter(e):
    e.widget['background'] = '#ddd'

def on_leave(e):
    e.widget['background'] = 'SystemButtonFace'

def center(win):
    """
    centers a tkinter window
    :param win: the main window or Toplevel window to center
    """
    win.update_idletasks()
    width = win.winfo_width()
    frm_width = win.winfo_rootx() - win.winfo_x()
    win_width = width + 2 * frm_width
    height = win.winfo_height()
    titlebar_height = win.winfo_rooty() - win.winfo_y()
    win_height = height + titlebar_height + frm_width
    x = win.winfo_screenwidth() // 2 - win_width // 2
    y = win.winfo_screenheight() // 2 - win_height // 2
    win.geometry('{}x{}+{}+{}'.format(width, height, x, y))
    win.deiconify()

if __name__ == '__main__':
    app = App()
    app.mainloop()